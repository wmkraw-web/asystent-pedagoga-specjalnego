import requests
import io
import streamlit as st
import json
import base64  # Dodana nowa biblioteka do odkodowywania obrazków!

try:
    import PyPDF2
except ImportError:
    pass
try:
    import docx
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    pass

# --- FUNKCJA: GENEROWANIE TEKSTU (GPT-4o-mini) ---
def call_openai_text(api_key, system_prompt, user_prompt, temperature=0.6):
    if not api_key:
        return "Błąd: Brak klucza API OpenAI."
    try:
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        payload = {
            "model": "gpt-4o-mini",
            "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
            "temperature": temperature
        }
        response = requests.post("https://api.openai.com/v1/chat/completions", headers=headers, json=payload, timeout=120)
        if response.ok:
            return response.json()["choices"][0]["message"]["content"]
        else:
            return f"Błąd API: {response.text}"
    except Exception as e:
        return f"Błąd komunikacji: {str(e)}"

# --- FUNKCJA: GENEROWANIE OBRAZU (SZYBKI MODEL GPT-IMAGE-1-MINI) ---
def call_openai_image(api_key, image_prompt):
    if not api_key:
        return None, "Brak klucza API."
    try:
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        payload = {
            "model": "gpt-image-1-mini",
            "prompt": image_prompt,
            "n": 1,
            "size": "1024x1024",
            "response_format": "b64_json" # Wymuszamy najszybszy format kodowania
        }
        response = requests.post("https://api.openai.com/v1/images/generations", headers=headers, json=payload, timeout=120)
        
        if response.ok:
            data = response.json()
            if "data" in data and len(data["data"]) > 0:
                img_data = data["data"][0]
                
                # SCENARIUSZ 1: Nowy model wysyła obraz w kodzie (b64_json) - NAJSZYBSZE
                if "b64_json" in img_data:
                    image_bytes = base64.b64decode(img_data["b64_json"])
                    return io.BytesIO(image_bytes), None
                    
                # SCENARIUSZ 2: Model wysyła klasyczny link (url)
                elif "url" in img_data:
                    image_url = img_data["url"]
                    img_response = requests.get(image_url)
                    if img_response.ok:
                        return io.BytesIO(img_response.content), None
                    else:
                        return None, "Błąd podczas pobierania wygenerowanego obrazka z linku."
                else:
                    return None, f"Błąd: Nieznany format odpowiedzi. Surowe dane: {json.dumps(data)}"
            else:
                return None, f"Błąd (Brak obrazka). Surowe dane z serwera: {json.dumps(data)}"
        else:
            try:
                err_data = response.json()
                return None, f"Odrzucono (Kod {response.status_code}). Powód: {json.dumps(err_data)}"
            except:
                return None, f"Błąd API Obrazów: {response.text}"
    except Exception as e:
        return None, f"Błąd komunikacji: {str(e)}"

# --- FUNKCJA: EKSPORT DO WORDA (.DOCX) ---
def create_word_document(title, content_text, image_bytes=None):
    doc = docx.Document()
    
    # Marginesy
    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(1)
    
    # Styl domyślny
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Nagłówek
    h = doc.add_heading(title.upper(), level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Notka o zgodności z MEN
    men_note = doc.add_paragraph("Dokument wygenerowany w oparciu o aktualne wytyczne MEN i podstawę programową.")
    men_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    men_note.runs[0].font.size = Pt(9)
    men_note.runs[0].font.italic = True
    
    doc.add_paragraph("-" * 80)
    
    # Jeśli mamy wygenerowany obrazek, wstawiamy go na początku Worda
    if image_bytes:
        doc.add_picture(image_bytes, width=Inches(4.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph() # Odstęp

    # Wstawianie tekstu (z obsługą pogrubień markdown)
    for line in content_text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        if line.startswith('### '): doc.add_heading(line[4:], level=3)
        elif line.startswith('## '): doc.add_heading(line[3:], level=2)
        elif line.startswith('# '): doc.add_heading(line[2:], level=1)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _add_bold_parts(p, line[2:])
        else:
            p = doc.add_paragraph()
            _add_bold_parts(p, line)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def _add_bold_parts(paragraph, text):
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        if i % 2 != 0: run.bold = True

# --- FUNKCJA: CZYTANIE PDF/DOCX ---
def extract_text_from_file(uploaded_file):
    text = ""
    try:
        if uploaded_file.name.endswith('.pdf'):
            reader = PyPDF2.PdfReader(uploaded_file)
            for page in reader.pages: text += page.extract_text() + "\n"
        elif uploaded_file.name.endswith('.docx'):
            doc = docx.Document(uploaded_file)
            for para in doc.paragraphs: 
                if para.text.strip(): text += para.text + "\n"
    except Exception:
        pass
    return text
